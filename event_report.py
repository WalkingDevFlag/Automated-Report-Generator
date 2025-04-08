# event_report.py
"""
Handles the generation of the 'Event Report' DOCX document.

Loads the specific template file ('event_report.docx'), replaces
text placeholders based on data from the Google Sheet row, and
inserts images into designated locations within the template.
"""
from __future__ import annotations
import os
import traceback
from docx import Document
from docx.shared import Inches
from docx.opc.exceptions import PackageNotFoundError

# Project Modules
import config # Import base configuration for paths and headers

# ===========================================================
# --- Event Report Specific Configuration ---
# ===========================================================

# Filename of the DOCX template located in the TEMPLATE_FOLDER
EVENT_REPORT_TEMPLATE_FILENAME = 'event_report.docx'

# Mapping from placeholder strings in the DOCX template (keys)
# to the corresponding configuration constants (values), which in turn
# map to the Google Sheet column headers.
EVENT_REPORT_PLACEHOLDERS = {
    '{{EventTitle}}': config.HEADER_EVENT_TITLE,
    '{{EventLocation}}': config.HEADER_EVENT_LOCATION,
    '{{EventIntroduction}}': config.HEADER_EVENT_INTRO,
    '{{EventObjective}}': config.HEADER_EVENT_OBJECTIVE,
    '{{EventDetails}}': config.HEADER_EVENT_DETAILS,
    '{{StartDate}}': config.HEADER_EVENT_START_DATE,
    '{{EndDate}}': config.HEADER_EVENT_END_DATE,
    # Note: Image placeholders below are just text markers for insertion points,
    # the text itself is removed during processing.
}

# --- Image Placeholder Tags ---
# These tags MUST EXACTLY match the plain text used in your .docx template
# where you want the respective images to be inserted.
IMAGE_PLACEHOLDER_TAG_1 = '{{EventBrochure}}' # Placeholder for the first image
IMAGE_PLACEHOLDER_TAG_2 = '{{EventImages}}'   # Placeholder for the second image

# --- Image Display Properties ---
IMAGE_WIDTH_INCHES = 5.0 # Default width for inserted images

# ===========================================================
# --- End Event Report Configuration ---
# ===========================================================


# --- Absolute Path for Template ---
EVENT_REPORT_TEMPLATE_PATH = os.path.join(config.TEMPLATE_FOLDER, EVENT_REPORT_TEMPLATE_FILENAME)


def _insert_image(document: Document, image_path: str | None, target_para_idx: int, image_log_name: str):
    """
    Helper function to insert an image into the document.

    Attempts to insert at a specific paragraph index (where a placeholder tag was found).
    If the index is invalid or insertion fails, it falls back to adding the image
    at the end of the document.

    Args:
        document (Document): The python-docx Document object to modify.
        image_path (str | None): The local file path to the image. If None or invalid, nothing happens.
        target_para_idx (int): The index of the paragraph where the image placeholder was found (-1 if not found).
        image_log_name (str): A descriptive name for logging purposes (e.g., "Image 1 (Brochure)").
    """
    if not image_path or not os.path.exists(image_path):
        # Only log a warning if a path was expected but is invalid/missing.
        if image_path:
            print(f"[WARNING] {image_log_name}: File path invalid or file not found, skipping insertion: {image_path}")
        return # Do nothing if no valid path

    image_inserted_successfully = False
    base_image_name = os.path.basename(image_path)

    # --- Attempt insertion at the target paragraph ---
    if target_para_idx != -1:
        try:
            # Check if the target index is still valid within the document's paragraphs
            if 0 <= target_para_idx < len(document.paragraphs):
                target_para = document.paragraphs[target_para_idx]
                print(f"[INFO] Attempting to add {image_log_name} ('{base_image_name}') into paragraph {target_para_idx}...")
                # Add the picture to a new run within the target paragraph
                # The placeholder text should have already been removed from this paragraph.
                run = target_para.add_run()
                # You could add a newline before the image if desired: target_para.add_run("\n")
                run.add_picture(image_path, width=Inches(IMAGE_WIDTH_INCHES))
                image_inserted_successfully = True
                print(f"[SUCCESS] Added {image_log_name} to paragraph {target_para_idx}.")
            else:
                # This might happen if paragraphs were deleted before image insertion
                print(f"[WARNING] Target paragraph index {target_para_idx} for {image_log_name} is out of bounds ({len(document.paragraphs)} paragraphs exist).")
        except FileNotFoundError: # Pillow/docx raises this if the path is bad *during* add_picture
             print(f"[ERROR] {image_log_name}: File not found during insertion attempt: {image_path}")
        except Exception as e:
            print(f"[ERROR] Failed to add {image_log_name} ('{base_image_name}') to target paragraph {target_para_idx}: {e}")
            if "unsupported image format" in str(e).lower() or isinstance(e, KeyError):
                 print("  >>> Suggestion: Ensure image format is supported by python-docx (PNG, JPG, GIF, BMP, TIFF) or was converted correctly.")
            # traceback.print_exc() # Uncomment for detailed debugging if needed

    # --- Fallback: Add image at the end of the document ---
    if not image_inserted_successfully:
        print(f"[INFO] Adding {image_log_name} ('{base_image_name}') at the end of the document (fallback)...")
        try:
            # This adds the image as a new paragraph at the end.
            document.add_picture(image_path, width=Inches(IMAGE_WIDTH_INCHES))
            print(f"[SUCCESS] Added {image_log_name} at the end.")
        except FileNotFoundError:
             print(f"[ERROR] {image_log_name}: File not found during fallback insertion attempt: {image_path}")
        except Exception as e:
            print(f"[ERROR] Failed to add {image_log_name} ('{base_image_name}') at the end: {e}")
            if "unsupported image format" in str(e).lower() or isinstance(e, KeyError):
                 print("  >>> Suggestion: Ensure image format is supported or was converted correctly.")
            # traceback.print_exc() # Uncomment for detailed debugging if needed


def generate_report(data: dict, image1_local_path: str | None = None, image2_local_path: str | None = None) -> Document | None:
    """
    Generates the Event Report DOCX document by filling a template.

    Loads the 'event_report.docx' template, replaces text placeholders defined
    in EVENT_REPORT_PLACEHOLDERS with values from the 'data' dictionary,
    finds image placeholder tags (IMAGE_PLACEHOLDER_TAG_1, IMAGE_PLACEHOLDER_TAG_2),
    removes the tags, and inserts the provided images at those locations.

    Args:
        data (dict): Dictionary mapping Google Sheet headers (constants from config.py)
                     to the corresponding cell values for the current row.
        image1_local_path (str | None): Local file path to the first image (Brochure).
        image2_local_path (str | None): Local file path to the second image (Event Images).

    Returns:
        docx.Document | None: The modified python-docx Document object ready to be saved,
                              or None if a critical error occurred (e.g., template not found).
    """
    print("[INFO] Starting Event Report generation...")

    # --- Load Template ---
    if not os.path.exists(EVENT_REPORT_TEMPLATE_PATH):
        print(f"[ERROR] Event Report template not found at: {EVENT_REPORT_TEMPLATE_PATH}")
        return None
    try:
        print(f"[INFO] Loading template: {EVENT_REPORT_TEMPLATE_PATH}")
        document = Document(EVENT_REPORT_TEMPLATE_PATH)
        print("[SUCCESS] Template loaded successfully.")
    except PackageNotFoundError:
         print(f"[ERROR] Failed to load template. File might be corrupted or not a valid DOCX: {EVENT_REPORT_TEMPLATE_PATH}")
         return None
    except Exception as e:
        print(f"[ERROR] Unexpected error loading template {EVENT_REPORT_TEMPLATE_PATH}: {e}")
        traceback.print_exc()
        return None

    # --- Placeholder Processing ---
    # Store paragraph indices where image tags are found
    image1_placeholder_para_idx = -1
    image2_placeholder_para_idx = -1
    found_tag1 = False # Debug flag
    found_tag2 = False # Debug flag

    print("[INFO] Processing text and image placeholders in template paragraphs...")
    # Iterate through paragraphs to replace text and find image locations
    for i, paragraph in enumerate(document.paragraphs):
        original_text = paragraph.text
        modified_text = original_text # Start with original text for modification
        text_was_modified = False
        is_image_para = False

        # --- Identify and Remove Image Placeholder Tags ---
        # Check for image tags first, store their paragraph index, and remove the tag text.
        if IMAGE_PLACEHOLDER_TAG_1 in modified_text:
            image1_placeholder_para_idx = i
            modified_text = modified_text.replace(IMAGE_PLACEHOLDER_TAG_1, "")
            found_tag1 = True
            text_was_modified = True
            is_image_para = True
            print(f"[DEBUG] Found Image Placeholder 1 tag ('{IMAGE_PLACEHOLDER_TAG_1}') in paragraph {i}")

        if IMAGE_PLACEHOLDER_TAG_2 in modified_text:
            # Important: Use elif if tags can't be in the same paragraph, else use if. Assuming they can be separate.
            if i == image1_placeholder_para_idx: # If tag 2 is in the *same* paragraph as tag 1
                 print(f"[WARNING] Image Placeholder 2 tag found in the same paragraph ({i}) as Tag 1. Tag 2 insertion may overwrite Tag 1 if target indices match later.")
                 # Keep the same index for now, insertion logic will handle it
            image2_placeholder_para_idx = i
            modified_text = modified_text.replace(IMAGE_PLACEHOLDER_TAG_2, "") # Remove from potentially already modified text
            found_tag2 = True
            text_was_modified = True
            is_image_para = True
            print(f"[DEBUG] Found Image Placeholder 2 tag ('{IMAGE_PLACEHOLDER_TAG_2}') in paragraph {i}")

        # --- Replace Text Placeholders ---
        # Use the defined mapping (placeholder -> config constant -> sheet header)
        for placeholder, sheet_header in EVENT_REPORT_PLACEHOLDERS.items():
            if placeholder in modified_text:
                value_from_sheet = data.get(sheet_header, '') # Get value from response data dict
                # Replace placeholder with actual value (ensure it's a string)
                new_text = modified_text.replace(placeholder, str(value_from_sheet))
                if new_text != modified_text: # Check if replacement actually happened
                    modified_text = new_text
                    text_was_modified = True
                    print(f"[DEBUG] Replaced '{placeholder}' in paragraph {i}.")

        # --- Update Paragraph Content ---
        # Only modify the paragraph if text was replaced OR if it contained an image tag.
        # This avoids modifying paragraphs unnecessarily.
        if text_was_modified:
            # Clear the original paragraph content (all runs)
            paragraph.clear()
            # Add the modified text back in a single new run
            # This helps maintain the paragraph's original style but loses run-level formatting within the replaced text.
            # For preserving run-level formatting, a more complex run-by-run replacement function would be needed (like in remedial_report).
            paragraph.add_run(modified_text)

    # Report if image tags were expected but not found
    if not found_tag1: print(f"[WARNING] Image Placeholder 1 tag ('{IMAGE_PLACEHOLDER_TAG_1}') was not found in the template.")
    if not found_tag2: print(f"[WARNING] Image Placeholder 2 tag ('{IMAGE_PLACEHOLDER_TAG_2}') was not found in the template.")

    # --- Insert Images ---
    print("[INFO] Inserting images into designated locations...")
    # Insert Image 1 (Brochure)
    _insert_image(document, image1_local_path, image1_placeholder_para_idx, "Image 1 (Brochure)")

    # Insert Image 2 (Event Images) - Determine target paragraph for Image 2
    target_para_idx_img2 = -1
    if image2_placeholder_para_idx != -1:
        target_para_idx_img2 = image2_placeholder_para_idx # Prefer specific placeholder paragraph
        print(f"[INFO] Target paragraph for Image 2 set to {target_para_idx_img2} (found tag '{IMAGE_PLACEHOLDER_TAG_2}').")
    elif image1_placeholder_para_idx != -1:
        target_para_idx_img2 = image1_placeholder_para_idx # Fallback: put near image 1 if tag 2 wasn't found
        print(f"[INFO] Target paragraph for Image 2 set to {target_para_idx_img2} (near Image 1, as Tag 2 wasn't found).")
    else:
        # If neither tag was found, insertion will default to the end of the document inside _insert_image
        print("[INFO] No specific placeholder paragraph found for Image 2. It will be added at the end if path is valid.")

    _insert_image(document, image2_local_path, target_para_idx_img2, "Image 2 (Event Images)")

    print("[SUCCESS] Event Report content generation complete.")
    return document