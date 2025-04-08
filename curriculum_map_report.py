# -*- coding: utf-8 -*-
# curriculum_map_report.py
"""
Handles the generation of the 'Curriculum Map Report' DOCX document.

Dynamically generates a Table of Contents page and a separate main body
page based on provided content, handling section inclusion/exclusion,
renumbering, and formatting.
"""

# --- Standard Library Imports ---
from __future__ import annotations
import os
import csv
import re
import traceback
import time

# --- Third-Party Library Imports ---
import docx # Import the main library namespace
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.exceptions import PackageNotFoundError # Specific exception for loading
from docx.text.paragraph import Paragraph # For type hinting
from docx.text.run import Run # For type hinting
from docx.table import Table, _Row, _Cell # For type hinting

# --- Project-Specific Imports ---
import config

# --- Import Helper Functions ---
HELPERS_AVAILABLE = True
try:
    from remedial_report import replace_text_preserving_formatting, find_paragraph_with_text
    print("[INFO][CurriculumMap] Successfully imported helpers from remedial_report.")
except ImportError as e:
    print(f"[ERROR][CurriculumMap] Could not import helpers from remedial_report.py: {e}")
    HELPERS_AVAILABLE = False
    # Define basic fallbacks if import fails
    def replace_text_preserving_formatting(para: Paragraph, ph: str, val: str):
        if ph in para.text: para.text = para.text.replace(ph, str(val) if val is not None else '')
    def find_paragraph_with_text(doc: docx.Document | object, txt: str) -> Paragraph | None: # Use docx.Document here
        if isinstance(doc, docx.Document): # Use docx.Document here
            for p in doc.paragraphs:
                if txt in p.text: return p
        return None

# ===========================================================
# --- Module Configuration ---
# ===========================================================
CURRICULUM_MAP_REPORT_TEMPLATE_FILENAME = 'curriculum_map_report.docx'
CURRICULUM_MAP_REPORT_TEMPLATE_PATH = os.path.join(config.TEMPLATE_FOLDER, CURRICULUM_MAP_REPORT_TEMPLATE_FILENAME)
TEXT_PLACEHOLDERS = {
    '{{FACULTY}}': config.HEADER_CM_FACULTY, '{{NAME OF SCHOOL}}': config.HEADER_CM_SCHOOL,
    '{{NAME OF DEPARTMENT}}': config.HEADER_CM_DEPARTMENT, '{{NAME OF EVENT}}': config.HEADER_CM_EVENT_NAME,
    '{{Nature of Event}}': config.HEADER_CM_NATURE, '{{Date of Event}}': config.HEADER_CM_DATE,
    '{{EventNameFooter}}': config.HEADER_CM_EVENT_NAME,
}
TABLE_PLACEHOLDERS_MAP = {
    '{{guest_table}}': 'guests_csv', '{{program_table}}': 'program_csv',
    '{{lectureplan_table}}': 'lecture_plan_csv', '{{schedule_table}}': 'schedule_csv',
    '{{attendance_table}}': 'attendance_csv',
}
IMAGE_PLACEHOLDERS_MAP = {
    '{{ClubLogo}}': 'logo', '{{photographs_placeholder}}': 'photos',
    '{{brochure_placeholder}}': 'brochure', '{{news_placeholder}}': 'news',
}
REPORT_SECTIONS = [ # Titles MUST match BODY headings
    {'index': 1, 'title': 'Introduction of the Event', 'text_placeholder': '{{introduction}}', 'data_key': config.HEADER_CM_INTRODUCTION},
    {'index': 2, 'title': 'Objective of the Event', 'text_placeholder': '{{objective}}', 'data_key': config.HEADER_CM_OBJECTIVE},
    {'index': 3, 'title': 'Beneficiaries of the Event', 'text_placeholder': '{{beneficiaries}}', 'data_key': config.HEADER_CM_BENEFICIARIES},
    {'index': 4, 'title': 'Details of the Guests', 'table_placeholder': '{{guest_table}}', 'asset_key': 'guests_csv'},
    {'index': 5, 'title': 'Brief Description of the Event', 'text_placeholder': '{{description}}', 'data_key': config.HEADER_CM_DESCRIPTION},
    {'index': 6, 'title': 'Program Scheme', 'table_placeholder': '{{program_table}}', 'asset_key': 'program_csv'},
    {'index': 7, 'title': 'Lecture Plan', 'table_placeholder': '{{lectureplan_table}}', 'asset_key': 'lecture_plan_csv'},
    {'index': 8, 'title': 'Photographs', 'image_placeholder': '{{photographs_placeholder}}', 'asset_key': 'photos'},
    {'index': 9, 'title': 'Brochure or Creative of the Event', 'image_placeholder': '{{brochure_placeholder}}', 'asset_key': 'brochure'},
    {'index': 10, 'title': 'Schedule of the Event', 'table_placeholder': '{{schedule_table}}', 'asset_key': 'schedule_csv'},
    {'index': 11, 'title': 'Attendance of the Event', 'table_placeholder': '{{attendance_table}}', 'asset_key': 'attendance_csv'},
    {'index': 12, 'title': 'News Publication', 'image_placeholder': '{{news_placeholder}}', 'asset_key': 'news'},
    {'index': 13, 'title': 'Feedback report of the Event', 'text_placeholder': '{{feedback}}', 'data_key': config.HEADER_CM_FEEDBACK},
    {'index': 14, 'title': 'Link of MUJ website stating that event is uploaded on Website', 'text_placeholder': '{{link}}', 'data_key': config.HEADER_CM_LINK},
]
CURRICULUM_MAP_TABLE_STYLE = 'Table Grid'
HEADER_ROW_FONT_BOLD = True
DOCUMENT_FONT_NAME = "Times New Roman"
FALLBACK_BORDER = {"sz": 4, "val": "single", "color": "000000", "space": "0"}

# ===========================================================
# --- Helper Functions ---
# ===========================================================

def set_table_borders(table: Table, border_settings: dict = FALLBACK_BORDER):
    """Applies basic black borders to all cells in a table."""
    try:
        tbl = table._tbl; tblPr = tbl.tblPr
        if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None: tblBorders = OxmlElement("w:tblBorders"); tblPr.append(tblBorders)
        border_keys = ["top", "left", "bottom", "right", "insideH", "insideV"]
        for key in border_keys:
            border_element = tblBorders.find(qn(f"w:{key}"))
            if border_element is None: border_element = OxmlElement(f"w:{key}"); tblBorders.append(border_element)
            for attr, value in border_settings.items(): border_element.set(qn(f"w:{attr}"), str(value))
    except AttributeError: print(f"  [ERROR][CurriculumMap] AttributeError during set_table_borders."); traceback.print_exc()
    except Exception as e: print(f"  [WARNING][CurriculumMap] Failed applying enforced table borders: {e}")

def _insert_table_from_csv_local(document: docx.Document, target_para: Paragraph | None, csv_file_path: str | None, table_log_id: str = "", table_style_name: str | None = None, header_bold: bool = True, apply_enforced_borders: bool = True) -> bool:
    """Reads CSV data and inserts it as a DOCX table immediately after the target_para."""
    log_prefix = f"  Table {table_log_id}: "
    if not target_para: print(f"{log_prefix}[ERROR] Target paragraph for insertion not found."); return False
    if not csv_file_path: print(f"{log_prefix}[INFO] No CSV path provided. Skipping."); return True
    if not os.path.exists(csv_file_path): print(f"{log_prefix}[ERROR] CSV not found: {csv_file_path}"); return False
    data = [];
    try:
        with open(csv_file_path, mode='r', newline='', encoding='utf-8-sig') as f:
            reader = csv.reader(f); data = [r for r in reader if any(c.strip() for c in r)]
        if not data: print(f"{log_prefix}[WARNING] CSV is empty. Skipping table creation."); return True
    except Exception as e: print(f"{log_prefix}[ERROR] Reading CSV failed: {e}"); return False
    num_rows=len(data); num_cols=max(len(r) for r in data) if data else 0
    if num_cols==0: print(f"{log_prefix}[WARNING] No columns in CSV. Skipping."); return True
    try:
        table=document.add_table(rows=num_rows, cols=num_cols); table.autofit=True
        style_applied=False
        if table_style_name:
            try: table.style=table_style_name; style_applied=True
            except Exception as style_e: print(f"{log_prefix}[WARNING] Applying style '{table_style_name}' failed: {style_e}.")
        if apply_enforced_borders: set_table_borders(table)
        for i, row_data in enumerate(data):
            if i < len(table.rows):
                cells=table.rows[i].cells
                for j, cell_text in enumerate(row_data):
                    if j < len(cells):
                        cp=cells[j].paragraphs[0];
                        for r in cp.runs: r.clear()
                        run=cp.add_run(str(cell_text).strip())
                        if i==0 and header_bold: run.font.bold=True
        target_para._p.addnext(table._tbl); return True
    except Exception as e: print(f"{log_prefix}[ERROR] Table creation/insertion failed: {e}"); traceback.print_exc(); return False

def set_run_font(run: Run, font_name: str):
    try:
        run.font.name = font_name; r = run._r; rPr = r.get_or_add_rPr()
        rFonts = rPr.first_child_found_in(qn("w:rFonts"))
        if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), font_name); rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name); rFonts.set(qn('w:cs'), font_name)
    except Exception as e: print(f"  [WARNING][CurriculumMap] Failed setting font '{font_name}': {e}")

def enforce_font_entire_document(document: docx.Document, font_name: str):
    print(f"  [INFO] Enforcing font '{font_name}'...")
    runs_processed = 0
    try:
        for para in document.paragraphs:
            for run in para.runs: set_run_font(run, font_name); runs_processed += 1
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs: set_run_font(run, font_name); runs_processed += 1
        for section in document.sections:
            for hf_part in [section.header, section.first_page_header, section.even_page_header,
                            section.footer, section.first_page_footer, section.even_page_footer]:
                if hf_part is not None:
                    for para in hf_part.paragraphs:
                        for run in para.runs: set_run_font(run, font_name); runs_processed += 1
                    for table in hf_part.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    for run in para.runs: set_run_font(run, font_name); runs_processed += 1
    except Exception as e: print(f"  [ERROR][CurriculumMap] Font enforcement failed: {e}")

def replace_header_footer_placeholders(document: docx.Document, placeholders_map: dict, data: dict):
    if not HELPERS_AVAILABLE: return
    try:
        for section in document.sections:
            for hf_part in [section.header, section.first_page_header, section.even_page_header,
                            section.footer, section.first_page_footer, section.even_page_footer]:
                if hf_part is not None:
                    for para in hf_part.paragraphs:
                        for ph, key in placeholders_map.items():
                            if ph in para.text: val = data.get(key, ''); replace_text_preserving_formatting(para, ph, val)
    except Exception as e: print(f"  [ERROR][CurriculumMap] H/F placeholder replacement failed: {e}")

def delete_paragraph(paragraph: Paragraph | None) -> bool:
    if not paragraph: return False
    try:
        p_element = paragraph._element; parent = p_element.getparent()
        if parent is not None: parent.remove(p_element); return True
        else: return False
    except Exception as e: print(f"  [WARNING][CurriculumMap] Error deleting paragraph: {e}"); return False

def delete_table(table: Table | None) -> bool:
    if not table: return False
    try:
        tbl_element = table._element; parent = tbl_element.getparent()
        if parent is not None: parent.remove(tbl_element); return True
        else: return False
    except Exception as e: print(f"  [WARNING][CurriculumMap] Error deleting table: {e}"); return False

# ===========================================================
# --- Main Report Generation Function ---
# ===========================================================

def generate_report(data: dict, assets: dict) -> docx.Document | None:
    """
    Generates the Curriculum Map Report DOCX document.
    """
    print("\n=== Starting Curriculum Map Report Generation ===")
    start_time = time.time()
    if not HELPERS_AVAILABLE:
        print("  [FATAL ERROR][CurriculumMap] Core helper functions missing.")
        return None

    # --- Load Template ---
    if not os.path.exists(CURRICULUM_MAP_REPORT_TEMPLATE_PATH):
        print(f"  [FATAL ERROR][CurriculumMap] Template not found: {CURRICULUM_MAP_REPORT_TEMPLATE_PATH}")
        return None
    try:
        print(f"  [INFO] Loading template: {CURRICULUM_MAP_REPORT_TEMPLATE_PATH}")
        # Explicitly use docx.Document to resolve potential NameError
        document = docx.Document(CURRICULUM_MAP_REPORT_TEMPLATE_PATH) # <<< CORRECTED LINE
        print("  [SUCCESS] Template loaded.")
    except PackageNotFoundError:
         print(f"  [FATAL ERROR][CurriculumMap] Failed loading template (corrupted/invalid format?): {CURRICULUM_MAP_REPORT_TEMPLATE_PATH}")
         return None
    except Exception as e:
        print(f"  [FATAL ERROR][CurriculumMap] Error loading template: {e}")
        traceback.print_exc()
        return None

    # --- Phase 1: Static Content ---
    print("\n  [INFO] Phase 1: Processing Static Content...")
    replace_header_footer_placeholders(document, TEXT_PLACEHOLDERS, data)
    # Insert Header Logo
    logo_placeholder_tag = '{{ClubLogo}}'; logo_asset_key = IMAGE_PLACEHOLDERS_MAP.get(logo_placeholder_tag)
    if logo_asset_key:
        logo_path = assets.get(logo_asset_key)
        if logo_path and os.path.exists(logo_path):
            logo_width = getattr(config, 'LOGO_WIDTH_INCHES', 1.0); inserted = False
            for section in document.sections:
                for header_part in [section.header, section.first_page_header, section.even_page_header]:
                    if not header_part: continue
                    logo_para = next((p for p in header_part.paragraphs if logo_placeholder_tag in p.text), None)
                    if logo_para:
                        try:
                            replace_text_preserving_formatting(logo_para, logo_placeholder_tag, '')
                            run = logo_para.add_run(); run.add_picture(logo_path, width=Inches(logo_width))
                            print(f"    [SUCCESS] Inserted logo into header.")
                            inserted = True; break
                        except Exception as e_logo: print(f"    [ERROR] Failed inserting logo: {e_logo}")
                if inserted: break

    # Replace Page 1 Placeholders
    page1_phs_keys = ['{{FACULTY}}', '{{NAME OF SCHOOL}}', '{{NAME OF DEPARTMENT}}', '{{NAME OF EVENT}}', '{{Nature of Event}}', '{{Date of Event}}']
    for para in document.paragraphs[:25]:
        for ph in page1_phs_keys:
            if ph in para.text:
                key = TEXT_PLACEHOLDERS.get(ph)
                if key: val = data.get(key, ''); replace_text_preserving_formatting(para, ph, val)

    # --- Phase 2: Populate Dynamic Sections & Track Presence ---
    print("\n  [INFO] Phase 2: Populating Dynamic Sections...")
    present_sections = []; placeholder_para_cache = {}
    for section_info in REPORT_SECTIONS:
        sec_idx = section_info['index']; sec_title = section_info['title']
        has_content = False;
        # Check/Populate Text
        if 'text_placeholder' in section_info and 'data_key' in section_info:
            placeholder = section_info['text_placeholder']; value = data.get(section_info['data_key'], '').strip()
            if value:
                target_para = placeholder_para_cache.get(placeholder) or find_paragraph_with_text(document, placeholder)
                if target_para: placeholder_para_cache[placeholder] = target_para; replace_text_preserving_formatting(target_para, placeholder, value); has_content = True
        # Check/Insert Table
        if 'table_placeholder' in section_info and 'asset_key' in section_info:
            placeholder = section_info['table_placeholder']; asset_path = assets.get(section_info['asset_key'])
            if asset_path and os.path.exists(asset_path):
                target_para = placeholder_para_cache.get(placeholder) or find_paragraph_with_text(document, placeholder)
                if target_para:
                    placeholder_para_cache[placeholder] = target_para
                    inserted = _insert_table_from_csv_local(document, target_para, asset_path, f"{sec_idx}", CURRICULUM_MAP_TABLE_STYLE, HEADER_ROW_FONT_BOLD, True)
                    if inserted: target_para.text = ""; has_content = True
        # Check/Insert Image
        if 'image_placeholder' in section_info and 'asset_key' in section_info:
            placeholder = section_info['image_placeholder']; asset_path = assets.get(section_info['asset_key'])
            if asset_path and os.path.exists(asset_path):
                target_para = placeholder_para_cache.get(placeholder) or find_paragraph_with_text(document, placeholder)
                if target_para:
                    placeholder_para_cache[placeholder] = target_para
                    _, ext = os.path.splitext(asset_path); is_image = ext.lower() in ['.png','.jpg','.jpeg','.gif','.bmp','.tiff']
                    should_insert = is_image or (section_info['asset_key'] != 'news')
                    if section_info['asset_key'] == 'news' and not is_image:
                        replace_text_preserving_formatting(target_para, placeholder, ''); has_content = True
                    elif should_insert:
                        try:
                            replace_text_preserving_formatting(target_para, placeholder, ''); run = target_para.add_run()
                            img_width = getattr(config, 'IMAGE_WIDTH_INCHES', 5.0); run.add_picture(asset_path, width=Inches(img_width)); has_content = True
                        except Exception as e_img: print(f"    [ERROR] Failed inserting image '{os.path.basename(asset_path)}': {e_img}")
        if has_content: present_sections.append(section_info)
    print(f"  [INFO] Identified {len(present_sections)} sections with content.")

    # --- Phase 3: Identify Content for Deletion ---
    print("\n  [INFO] Phase 3: Identifying Content for Deletion...")
    toc_heading_text = "Content of Report"
    toc_start_para_idx = -1
    toc_paras_to_delete = []
    body_elements_to_delete = []
    last_toc_para_idx = -1 # Initialize

    toc_start_para = next((p for i, p in enumerate(document.paragraphs) if toc_heading_text in p.text), None)
    if toc_start_para:
        try: toc_start_para_idx = document.paragraphs.index(toc_start_para)
        except ValueError: toc_start_para_idx = -1

    if toc_start_para_idx != -1:
        # Find ToC lines to delete
        for i in range(toc_start_para_idx + 1, len(document.paragraphs)):
            para = document.paragraphs[i]; text = para.text.strip()
            match = re.match(r"^\s*(?:\{\{(\d+)\.\}\}|(\d+)\.)\s*(.*)", text)
            if match:
                last_toc_para_idx = i
                original_idx = int(match.group(1) or match.group(2)); is_present = any(s['index'] == original_idx for s in present_sections)
                if not is_present: toc_paras_to_delete.append(para)
            elif text: break
        if last_toc_para_idx == -1: last_toc_para_idx = toc_start_para_idx

        # Find Body elements to delete (Refined Logic)
        sections_to_delete = [s for s in REPORT_SECTIONS if s['index'] not in [p['index'] for p in present_sections]]
        if sections_to_delete:
             # print(f"    Identifying body content for sections to delete: {[s['index'] for s in sections_to_delete]}")
             body = document.element.body
             element_index = 0
             while element_index < len(body):
                 element = body[element_index]
                 element_index += 1
                 para_obj = None; heading_match = None; is_target_heading = False; target_section_idx = -1

                 if isinstance(element, docx.oxml.text.paragraph.CT_P):
                      para_obj = Paragraph(element, document); para_text = para_obj.text.strip()
                      heading_match = re.match(r"^\s*(\d+)\s*\.\s*(.*)", para_text)
                      if heading_match:
                           heading_idx = int(heading_match.group(1)); heading_title = heading_match.group(2).strip()
                           section_def = next((s for s in sections_to_delete if s['index'] == heading_idx and s['title'].strip().lower() == heading_title.lower()), None)
                           if section_def: is_target_heading = True; target_section_idx = heading_idx

                 if is_target_heading:
                      body_elements_to_delete.append(element)
                      content_element_index = element_index
                      while content_element_index < len(body):
                          content_element = body[content_element_index]; content_para = None; is_next_heading = False
                          if isinstance(content_element, docx.oxml.text.paragraph.CT_P):
                              content_para = Paragraph(content_element, document)
                              if re.match(r"^\s*(\d+)\s*\.\s*(.*)", content_para.text.strip()): is_next_heading = True
                          if is_next_heading: break
                          else: body_elements_to_delete.append(content_element); content_element_index += 1
                      element_index = content_element_index # Adjust main loop index
    else:
        print(f"  [WARNING] ToC heading '{toc_heading_text}' not found. Cleanup skipped.")

    # --- Phase 4: Perform Deletions ---
    print("\n  [INFO] Phase 4: Performing Deletions...")
    deleted_body_count = 0
    if body_elements_to_delete:
        body = document.element.body; elements_to_delete_set = set(body_elements_to_delete)
        for element in list(body):
            if element in elements_to_delete_set:
                try: body.remove(element); deleted_body_count += 1
                except ValueError: pass
        print(f"    Deleted {deleted_body_count} body elements.")
    deleted_toc_count = 0
    if toc_paras_to_delete:
        for para in toc_paras_to_delete:
             if delete_paragraph(para): deleted_toc_count += 1
        print(f"    Deleted {deleted_toc_count} ToC lines.")
        # Re-calculate last_toc_para_idx
        if toc_start_para_idx != -1:
             temp_last_idx = -1
             for i in range(toc_start_para_idx + 1, len(document.paragraphs)):
                 para = document.paragraphs[i]; text = para.text.strip()
                 match = re.match(r"^\s*(\d+)\.\s*(.*)", text)
                 if match: temp_last_idx = i
                 elif text: break
             last_toc_para_idx = temp_last_idx if temp_last_idx != -1 else toc_start_para_idx

    # --- Phase 5: Insert Page Break ---
    print("\n  [INFO] Phase 5: Inserting Page Break...")
    page_break_inserted = False
    if last_toc_para_idx != -1 and 0 <= last_toc_para_idx < len(document.paragraphs):
        try:
            para_to_insert_after = document.paragraphs[last_toc_para_idx]
            new_para = para_to_insert_after.insert_paragraph_after("")
            run = new_para.add_run(); run.add_break(WD_BREAK.PAGE); page_break_inserted = True
            print("    Page break inserted after last ToC item.")
        except Exception as pb_e: print(f"    [WARNING] Failed inserting page break after ToC item: {pb_e}")
    elif toc_start_para:
        try:
            new_para = toc_start_para.insert_paragraph_after("")
            run = new_para.add_run(); run.add_break(WD_BREAK.PAGE); page_break_inserted = True
            print("    Page break inserted after ToC title (Fallback).")
        except Exception as pb_e_fb: print(f"    [WARNING] Failed inserting page break via fallback: {pb_e_fb}")
    if not page_break_inserted: print("    [WARNING] Could not insert page break.")

    # --- Phase 6: Renumber ToC and Headings ---
    print("\n  [INFO] Phase 6: Renumbering...")
    renumbered_toc_count = 0; page_break_para = None
    if page_break_inserted:
         try: page_break_para = next(p for p in document.paragraphs if p.runs and p.runs[-1].element.xpath('.//w:br[@w:type="page"]'))
         except StopIteration: pass
    page_break_idx = -1
    if page_break_para:
        try: page_break_idx = document.paragraphs.index(page_break_para)
        except ValueError: pass

    if toc_start_para_idx != -1:
        current_toc_num = 1
        toc_process_end_idx = page_break_idx if page_break_idx != -1 else len(document.paragraphs)
        for i in range(toc_start_para_idx + 1, toc_process_end_idx):
            if i >= len(document.paragraphs): break
            para = document.paragraphs[i]; match = re.match(r"^\s*(\d+)\.\s*(.*)", para.text.strip())
            if match:
                title_in_toc = match.group(2).strip()
                section_def = next((s for s in present_sections if s['title'].strip().lower() == title_in_toc.lower()), None)
                if section_def:
                    new_line = f"{current_toc_num}. {section_def['title']}"
                    para.clear(); run = para.add_run(new_line); run.font.bold = False
                    current_toc_num += 1; renumbered_toc_count += 1
        print(f"    Renumbered {renumbered_toc_count} ToC items.")

    # Renumber Body Headings
    current_heading_num = 1; renumbered_heading_count = 0
    start_heading_renumber_idx = page_break_idx + 1 if page_break_idx != -1 else 0
    for i in range(start_heading_renumber_idx, len(document.paragraphs)):
        para = document.paragraphs[i]
        match = re.match(r"^\s*(\d+)\s*\.\s*(.*)", para.text.strip())
        if match:
            title = match.group(2).strip()
            is_present = any(s['title'].strip().lower() == title.lower() for s in present_sections)
            if is_present:
                new_heading_text = f"{current_heading_num}. {title}"
                style = para.style
                para.clear(); run = para.add_run(new_heading_text); run.font.bold = True
                try: para.style = style
                except: pass
                current_heading_num += 1; renumbered_heading_count += 1
    print(f"    Renumbered {renumbered_heading_count} body headings.")

    # --- Phase 7: Final Font Enforcement ---
    print("\n  [INFO] Phase 7: Final Font Enforcement...")
    enforce_font_entire_document(document, DOCUMENT_FONT_NAME)

    # --- Completion ---
    total_time = time.time() - start_time
    print(f"\n[SUCCESS] Curriculum Map Report generation complete ({total_time:.2f}s).")
    print("=== Curriculum Map Report Generation Finished ===")
    return document