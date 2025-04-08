# remedial_report.py
"""
Handles the generation of the 'Remedial Class Report' DOCX document.

Loads the specific template file ('remedial_report.docx'), replaces
text placeholders, finds table placeholder tags, reads data from
provided local CSV files, and inserts formatted tables into the document
at the designated locations. Uses run-level text replacement to better
preserve formatting.
"""
from __future__ import annotations
import os
import csv
import copy # Needed for deep copying runs during text replacement
import traceback
from docx import Document
from docx.shared import Pt # Points for font size (example, not actively used unless styling)
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Example alignment
from docx.table import _Cell # Type hint example
from docx.text.paragraph import Paragraph # Type hint example
from docx.text.run import Run # Type hint example
from docx.oxml.shared import OxmlElement # For table borders if needed
from docx.oxml.ns import qn # For table borders if needed
from docx.opc.exceptions import PackageNotFoundError

# Project Modules
import config

# ===========================================================
# --- Remedial Class Report Specific Configuration ---
# ===========================================================

# Filename of the DOCX template located in the TEMPLATE_FOLDER
REMEDIAL_REPORT_TEMPLATE_FILENAME = 'remedial_report.docx'
REMEDIAL_REPORT_TEMPLATE_PATH = os.path.join(config.TEMPLATE_FOLDER, REMEDIAL_REPORT_TEMPLATE_FILENAME)

# --- Text Placeholder Mapping ---
# Maps DOCX placeholders (keys) to Sheet Header constants (values)
REMEDIAL_REPORT_PLACEHOLDERS = {
    '{{SchoolName}}': config.HEADER_REMEDIAL_SCHOOL,
    '{{DepartmentName}}': config.HEADER_REMEDIAL_DEPARTMENT,
    '{{ProgramName}}': config.HEADER_REMEDIAL_PROGRAM,
    '{{CourseName}}': config.HEADER_REMEDIAL_COURSE,
    '{{CourseCode}}': config.HEADER_REMEDIAL_COURSE_CODE,
    '{{CourseCredit}}': config.HEADER_REMEDIAL_COURSE_CREDIT,
    '{{TotalHoursDMS}}': config.HEADER_REMEDIAL_TOTAL_HOURS,
    '{{TotalRemedialClasses}}': config.HEADER_REMEDIAL_CLASSES_TAKEN,
    '{{FacultyName}}': config.HEADER_REMEDIAL_FACULTY,
    '{{Session}}': config.HEADER_REMEDIAL_SESSION,
    '{{Semester}}': config.HEADER_REMEDIAL_SEMESTER,
    '{{CriteriaText}}': config.HEADER_REMEDIAL_CRITERIA,
    # Note: Table placeholders {{table1}}, {{table2}}, {{table3}} are handled separately
}

# --- Table Placeholder Tags ---
# These tags MUST EXACTLY match the plain text in your .docx template
# where the corresponding CSV data should be inserted as a table.
# IMPORTANT: Each tag should be in its own, separate paragraph in the template.
TABLE_PLACEHOLDER_TAG_1 = '{{table1}}' # Placeholder for Student List/Performance CSV data
TABLE_PLACEHOLDER_TAG_2 = '{{table2}}' # Placeholder for Timetable CSV data
TABLE_PLACEHOLDER_TAG_3 = '{{table3}}' # Placeholder for Attendance CSV data

# --- Table Styling ---
# Recommend using a pre-defined style from your Word template for consistency.
# Ensure this style name exists in your remedial_report.docx template.
TABLE_STYLE = 'Table Grid'      # Example: Use Word's built-in 'Table Grid' style
HEADER_ROW_FONT_BOLD = True     # Make the first row (headers) bold
# FONT_NAME = "Calibri" # Deprecated: Styling should ideally be in the Word template's TABLE_STYLE

# ===========================================================
# --- Helper Functions ---
# ===========================================================

def replace_text_preserving_formatting(paragraph: Paragraph, placeholder: str, value: str):
    """
    Replaces placeholder text within a paragraph while preserving formatting.

    Searches for the placeholder, potentially spanning multiple runs.
    Replaces the placeholder text with the provided value, applying the
    formatting of the *first* run of the original placeholder to the
    entire inserted value.

    Args:
        paragraph (Paragraph): The python-docx Paragraph object to modify.
        placeholder (str): The placeholder text to find (e.g., '{{PlaceholderName}}').
        value (str): The text value to insert in place of the placeholder.
    """
    # Ensure value is a string, handle None gracefully
    str_value = str(value) if value is not None else ''
    placeholder_len = len(placeholder)
    start_marker = placeholder[:2] # Typically "{{"
    end_marker = placeholder[-2:]   # Typically "}}"

    # --- Phase 1: Find the runs involved ---
    runs = paragraph.runs
    start_run_idx, start_char_offset = -1, -1
    end_run_idx, end_char_offset = -1, -1
    current_match_str = ""
    runs_in_match = []

    # Iterate through runs to find the start and end of the placeholder
    for i, run in enumerate(runs):
        run_text = run.text
        for j, char in enumerate(run_text):
            current_match_str += char

            # Check if we found the start marker
            if start_run_idx == -1 and current_match_str.endswith(start_marker):
                start_run_idx = i
                # Calculate offset within the starting run
                start_char_offset = j - len(start_marker) + 1
                runs_in_match = [(i, run)] # Start tracking runs in the match

            # If we've found the start, keep adding runs until we find the end
            elif start_run_idx != -1:
                if i not in [r[0] for r in runs_in_match]:
                    runs_in_match.append((i, run))

                # Check if the full placeholder is found
                if current_match_str.endswith(placeholder):
                    end_run_idx = i
                    end_char_offset = j # Index of the last char of the placeholder
                    # Found the complete placeholder, break loops
                    # print(f"[DEBUG] Found '{placeholder}' spanning runs {start_run_idx} to {end_run_idx}")
                    goto_replacement = True
                    break # Break inner char loop
        if end_run_idx != -1:
            break # Break outer run loop

    # If placeholder wasn't fully found in this paragraph
    if start_run_idx == -1 or end_run_idx == -1:
        # print(f"[DEBUG] Placeholder '{placeholder}' not fully found in paragraph.")
        return # Nothing to replace

    # --- Phase 2: Replace the text ---
    # Keep the style of the very first character of the placeholder
    style_run = runs[start_run_idx]

    # Iterate through the runs identified as part of the placeholder
    for i, run in runs_in_match:
        original_run_text = run.text
        start_clear = 0
        end_clear = len(original_run_text)

        # Determine portion of the run's text to clear
        if i == start_run_idx:
            start_clear = start_char_offset
        if i == end_run_idx:
            end_clear = end_char_offset + 1 # Clear up to and including the last char

        # Build the new text for this specific run
        new_text_parts = []
        # Keep text *before* the placeholder part in this run
        if start_clear > 0:
            new_text_parts.append(original_run_text[:start_clear])

        # Insert the *entire* replacement value into the *first run only*
        if i == start_run_idx:
            new_text_parts.append(str_value)

        # Keep text *after* the placeholder part in this run
        if end_clear < len(original_run_text):
             new_text_parts.append(original_run_text[end_clear:])

        # Update the run's text
        run.text = "".join(new_text_parts)
        # print(f"[DEBUG] Run {i} updated. Original='{original_run_text}', New='{run.text}'")

    # --- Phase 3: Apply style to the inserted value ---
    # Apply the style of the original placeholder's first run to the *entire* first run
    # (which now contains the inserted value). This assumes the value should take
    # on the style of the placeholder it replaced.
    if start_run_idx < len(paragraph.runs):
        target_run = paragraph.runs[start_run_idx]
        # Copy basic formatting properties
        target_run.bold = style_run.bold
        target_run.italic = style_run.italic
        target_run.underline = style_run.underline
        target_run.font.name = style_run.font.name
        if style_run.font.size: target_run.font.size = style_run.font.size
        if style_run.font.color.rgb: target_run.font.color.rgb = style_run.font.color.rgb
        # Copy other properties as needed (e.g., highlight, strike)


def find_paragraph_with_text(document: Document, text_to_find: str) -> Paragraph | None:
    """
    Finds the first paragraph object containing specific text within the document body.

    Checks main paragraphs first, then checks inside table cells.

    Args:
        document (Document): The python-docx Document object to search within.
        text_to_find (str): The exact text string to locate.

    Returns:
        Paragraph | None: The first Paragraph object containing the text, or None if not found.
    """
    # Search main body paragraphs
    for i, para in enumerate(document.paragraphs):
        if text_to_find in para.text:
            # print(f"[DEBUG] Found '{text_to_find}' in main paragraph index {i}")
            return para

    # Search paragraphs within table cells
    for table_idx, table in enumerate(document.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                     if text_to_find in para.text:
                        print(f"[WARNING] Found placeholder '{text_to_find}' inside table cell (Table {table_idx}, Row {row_idx}, Col {col_idx}). Table insertion logic might require adjustment if this wasn't intended.")
                        # Returning the paragraph itself for replacement, but insertion 'after' might be tricky.
                        return para

    print(f"[WARNING] Placeholder text '{text_to_find}' not found anywhere in the document body or tables.")
    return None

def _insert_table_from_csv(document: Document, target_para: Paragraph | None, csv_file_path: str | None, table_num_for_log: str = "") -> bool:
    """
    Reads data from a CSV file and inserts it as a DOCX table after the target paragraph.

    Args:
        document (Document): The main python-docx Document object.
        target_para (Paragraph | None): The paragraph object AFTER which the table should be inserted.
                                        Typically the paragraph containing the placeholder tag (e.g., '{{table1}}').
        csv_file_path (str | None): The full path to the local CSV file.
        table_num_for_log (str): An identifier for logging (e.g., "1", "2", "3").

    Returns:
        bool: True if the table was successfully created and inserted (even if CSV was empty),
              False if an error occurred or inputs were invalid.
    """
    log_prefix = f"  Table {table_num_for_log}: " if table_num_for_log else "  "

    # --- Input Validation ---
    if not target_para:
        print(f"{log_prefix}[ERROR] Cannot insert table - target placeholder paragraph object was not provided (likely not found).")
        return False
    if not csv_file_path:
        print(f"{log_prefix}[WARNING] Cannot insert table - CSV file path was not provided. Skipping.")
        # Returning True because the 'placeholder' exists, but no data to insert. Allows clearing the placeholder.
        # Change to False if an empty path should be treated as a hard error.
        return True
    if not os.path.exists(csv_file_path):
        print(f"{log_prefix}[ERROR] Cannot insert table - CSV file not found at path: {csv_file_path}")
        return False

    print(f"{log_prefix}[INFO] Reading data from CSV: {os.path.basename(csv_file_path)}")
    data = []
    try:
        # Use utf-8-sig to handle potential Byte Order Mark (BOM) in CSV files
        with open(csv_file_path, mode='r', newline='', encoding='utf-8-sig') as infile:
            reader = csv.reader(infile)
            # Read all rows, skipping completely empty ones
            data = [row for row in reader if any(field.strip() for field in row)]
            if not data:
                 print(f"{log_prefix}[WARNING] No data rows found in CSV file.")
                 # Proceed to insert an empty table frame if desired.
            else:
                 print(f"{log_prefix}[INFO] Read {len(data)} data rows from CSV.")

    except Exception as e:
        print(f"{log_prefix}[ERROR] Failed to read CSV file {csv_file_path}: {e}")
        traceback.print_exc()
        return False

    # --- Determine Table Dimensions ---
    num_rows = len(data)
    # Determine num_cols safely, handle empty data or rows with varying lengths
    num_cols = 0
    if data:
        num_cols = max(len(row) for row in data) if data else 0

    if num_cols == 0 and num_rows > 0:
         print(f"{log_prefix}[WARNING] CSV contains rows but no columns could be determined (all fields empty?). Skipping table insertion.")
         return False # Cannot create a table with 0 columns
    # If num_rows is 0, num_cols will also be 0 here.

    # --- Create Table Structure ---
    # Add table requires at least 1 row/col if creating an empty frame.
    table_rows_to_create = max(1, num_rows)
    table_cols_to_create = max(1, num_cols) # Ensure at least 1 column if creating frame

    print(f"{log_prefix}[INFO] Creating table with {num_rows} data rows ({table_rows_to_create} total) and {num_cols} columns ({table_cols_to_create} total). Style: '{TABLE_STYLE}'")
    try:
        # Create the table object, applying the specified style
        table = document.add_table(rows=table_rows_to_create, cols=table_cols_to_create, style=TABLE_STYLE)
        table.autofit = True # Adjust column widths based on content

        # --- Populate Table Cells (if data exists) ---
        if data:
            print(f"{log_prefix}[INFO] Populating table cells...")
            for i, row_data in enumerate(data):
                if i < len(table.rows): # Ensure row index is valid
                    row_cells = table.rows[i].cells
                    for j, cell_text in enumerate(row_data):
                        if j < len(row_cells): # Ensure column index is valid
                            # Clean cell text and get the target cell paragraph
                            cleaned_text = str(cell_text).strip()
                            cell_paragraph = row_cells[j].paragraphs[0]
                            # Clear any default text before adding new run
                            for run in cell_paragraph.runs: run.clear()
                            # Add the text and apply formatting
                            run = cell_paragraph.add_run(cleaned_text)
                            if i == 0 and HEADER_ROW_FONT_BOLD:
                                run.font.bold = True
                            # cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT # Optional: Set cell alignment
            print(f"{log_prefix}[INFO] Table cell population complete.")
        else:
             print(f"{log_prefix}[INFO] CSV was empty. Added an empty table frame.")
             # Clear the default text from the minimum 1x1 cell if needed
             if table_rows_to_create == 1 and table_cols_to_create == 1:
                 table.cell(0,0).text = "" # Clear the single cell

        # --- Insert the table into the document ---
        # This is the crucial step: insert the table's XML element (_tbl)
        # immediately after the target paragraph's XML element (_p).
        print(f"{log_prefix}[INFO] Inserting table into document structure...")
        target_para._p.addnext(table._tbl)
        print(f"{log_prefix}[SUCCESS] Table inserted successfully.")
        return True

    except Exception as e:
        print(f"{log_prefix}[ERROR] Failed to create or insert table from {csv_file_path}: {e}")
        traceback.print_exc()
        return False

# ===========================================================
# --- Main Report Generation Function ---
# ===========================================================

def generate_report(data: dict, csv_file_path_1: str | None = None, csv_file_path_2: str | None = None, csv_file_path_3: str | None = None) -> Document | None:
    """
    Generates the Remedial Class Report DOCX document.

    Loads the template, replaces text placeholders using run-level formatting preservation,
    locates table placeholder tags, inserts tables from the provided CSV file paths,
    and clears the placeholder tags.

    Args:
        data (dict): Dictionary mapping Google Sheet headers (constants from config.py)
                     to cell values for the current row.
        csv_file_path_1 (str | None): Local path to the CSV for Table 1 (Student List).
        csv_file_path_2 (str | None): Local path to the CSV for Table 2 (Timetable).
        csv_file_path_3 (str | None): Local path to the CSV for Table 3 (Attendance).

    Returns:
        docx.Document | None: The modified python-docx Document object, or None if a
                              critical error occurred (e.g., template not found).
    """
    print("[INFO] Starting Remedial Class Report generation...")

    # --- Load Template ---
    if not os.path.exists(REMEDIAL_REPORT_TEMPLATE_PATH):
        print(f"[ERROR] Remedial Report template not found: {REMEDIAL_REPORT_TEMPLATE_PATH}")
        return None
    try:
        print(f"[INFO] Loading template: {REMEDIAL_REPORT_TEMPLATE_PATH}")
        document = Document(REMEDIAL_REPORT_TEMPLATE_PATH)
        print("[SUCCESS] Template loaded successfully.")
    except PackageNotFoundError:
         print(f"[ERROR] Failed to load template. File might be corrupted or not a valid DOCX: {REMEDIAL_REPORT_TEMPLATE_PATH}")
         return None
    except Exception as e:
        print(f"[ERROR] Unexpected error loading template {REMEDIAL_REPORT_TEMPLATE_PATH}: {e}")
        traceback.print_exc()
        return None

    # --- Text Replacement ---
    print("[INFO] Replacing text placeholders (preserving formatting)...")
    placeholders_found_count = 0
    # Replace in main body paragraphs
    for para in document.paragraphs:
        for placeholder, sheet_header in REMEDIAL_REPORT_PLACEHOLDERS.items():
            # Check if placeholder *might* be in paragraph before attempting complex replace
            if placeholder in para.text:
                value = data.get(sheet_header, '') # Default to empty string if header missing in data
                replace_text_preserving_formatting(para, placeholder, value)
                placeholders_found_count += 1 # Count attempts, not successes

    # Replace in table cells (if template has pre-existing tables with placeholders)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                     for placeholder, sheet_header in REMEDIAL_REPORT_PLACEHOLDERS.items():
                        if placeholder in para.text:
                            value = data.get(sheet_header, '')
                            replace_text_preserving_formatting(para, placeholder, value)
                            placeholders_found_count += 1
    print(f"[INFO] Text placeholder replacement attempts complete (attempted on {placeholders_found_count} potential finds).")


    # --- Table Insertion ---
    # IMPORTANT Template Structure Note:
    # For table insertion to work correctly:
    # 1. Ensure each heading (e.g., "2. List of identified students...") is in its OWN paragraph.
    # 2. Ensure the corresponding table placeholder tag (e.g., "{{table1}}") is ALSO in its OWN,
    #    SEPARATE paragraph immediately following the heading paragraph.
    print("\n[INFO] Processing table insertions from CSV files...")

    # Define table processing steps (tag, csv_path, log_id)
    table_tasks = [
        (TABLE_PLACEHOLDER_TAG_1, csv_file_path_1, "1 (Student List)"),
        (TABLE_PLACEHOLDER_TAG_2, csv_file_path_2, "2 (Timetable)"),
        (TABLE_PLACEHOLDER_TAG_3, csv_file_path_3, "3 (Attendance)"),
    ]

    # Process tables sequentially
    for tag, csv_path, log_id in table_tasks:
        print(f"\n[INFO] Processing {log_id} - Placeholder: '{tag}'")
        # Find the paragraph containing the placeholder tag *each time*,
        # as inserting previous tables might shift paragraph indices.
        para_placeholder = find_paragraph_with_text(document, tag)

        if para_placeholder:
            # Attempt to insert the table from the corresponding CSV file
            table_inserted = _insert_table_from_csv(document, para_placeholder, csv_path, log_id)

            # Clear the placeholder tag paragraph AFTER attempting insertion.
            # We clear it even if insertion "failed" due to empty CSV or minor error,
            # as long as the placeholder paragraph itself was found.
            # This prevents the tag {{tableX}} appearing in the final report.
            print(f"  Table {log_id}: Clearing placeholder paragraph text '{tag}'...")
            para_placeholder.text = ""
            # Optional: Remove the (now empty) placeholder paragraph entirely.
            # try:
            #     p_element = para_placeholder._element
            #     if p_element.getparent() is not None:
            #         p_element.getparent().remove(p_element)
            #     print(f"  Table {log_id}: Removed empty placeholder paragraph.")
            # except Exception as del_e:
            #     print(f"  Table {log_id}: [WARNING] Could not remove empty placeholder paragraph: {del_e}")

        else:
            # Log if the placeholder tag itself wasn't found in the template
            print(f"  Table {log_id}: [WARNING] Placeholder tag '{tag}' not found in the template. Skipping table insertion.")
            # No need to check csv_path if the tag isn't there.

    print("\n[SUCCESS] Remedial Class Report content generation complete.")
    return document