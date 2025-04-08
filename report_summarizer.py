# report_summarizer.py
"""
Handles the summarization of uploaded PDF or DOCX reports.

Extracts text and images from the input document, uses a pre-trained
transformer model (e.g., BART, DistilBART) via the Hugging Face
`transformers` library to generate a summary, and creates a new
DOCX document containing the summary text and extracted images.
"""
from __future__ import annotations
import os
import io
import re
import traceback
import time # For basic timing/diagnostics

# --- PDF Processing Library ---
try:
    import fitz # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    print("[ERROR] Required library 'PyMuPDF' not found.")
    print("  >>> Please install it: pip install PyMuPDF")
    print("  >>> PDF processing will be unavailable for summarizer.")
    PYMUPDF_AVAILABLE = False

# --- Hugging Face & ML Libraries ---
try:
    import torch
    import transformers
    from huggingface_hub import hf_hub_download
    from huggingface_hub.utils import EntryNotFoundError as HubNotFoundError
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    print("[ERROR] Required libraries ('transformers', 'torch', 'huggingface_hub') not found.")
    print("  >>> Please install them: pip install transformers torch huggingface_hub")
    print("  >>> Summarization feature will be unavailable.")
    TRANSFORMERS_AVAILABLE = False

# --- Document & Image Processing Libraries ---
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.opc.exceptions import PackageNotFoundError
    from docx.text.paragraph import Paragraph # Type hinting
    from docx.text.run import Run # Type hinting
    DOCX_AVAILABLE = True
except ImportError:
    print("[ERROR] Required library 'python-docx' not found.")
    print("  >>> Please install it: pip install python-docx")
    print("  >>> DOCX processing will be unavailable.")
    DOCX_AVAILABLE = False

try:
    from PIL import Image, ImageDraw # Pillow for image handling
    PIL_AVAILABLE = True
except ImportError:
    print("[WARNING] Library 'Pillow' not found (pip install Pillow).")
    print("  >>> Image saving/conversion might be limited.")
    PIL_AVAILABLE = False

# Standard library for network errors
import requests.exceptions

# Project Modules
import config # For TEMP_FOLDER etc.

# ===========================================================
# --- Summarizer Configuration ---
# ===========================================================
# Choose a summarization model from Hugging Face Hub
# Smaller/faster options: 'sshleifer/distilbart-cnn-6-6', 'google/pegasus-xsum'
# Larger/potentially better options: 'facebook/bart-large-cnn', 'google/pegasus-large'
SUMMARIZER_MODEL = "facebook/bart-large-cnn"
# SUMMARIZER_MODEL = "sshleifer/distilbart-cnn-12-6" # Previous choice

# --- Summarization Parameters ---
# Desired summary length as a fraction of the input text length (token count)
SUMMARY_MAX_LENGTH_FACTOR = 0.40 # Allow summary up to 40% of original length
SUMMARY_MIN_LENGTH_FACTOR = 0.15 # Require summary at least 15% of original length
# Absolute token limits to prevent excessively long/short summaries or model errors
SUMMARY_ABS_MAX_LENGTH = 1024 # Max tokens for the generated summary (adjust based on model)
SUMMARY_ABS_MIN_LENGTH = 50   # Min tokens for the generated summary
# Model input truncation - If input exceeds model max, should we truncate manually?
MANUAL_TRUNCATE_INPUT = True # True = truncate before pipeline; False = let pipeline handle it
MANUAL_TRUNCATE_TOKEN_LIMIT = 1000 # Max input tokens if MANUAL_TRUNCATE_INPUT=True (use < model max)

# --- Image Handling ---
IMAGE_WIDTH_INCHES_SUMMARY = 4.5    # Width for placing extracted images in the summary doc
TEMP_IMAGE_PREFIX = "extracted_img_" # Prefix for temporary image filenames
MIN_IMAGE_DIMENSION_PDF = 50        # Ignore tiny images extracted from PDFs (pixels)

# --- Debugging/Performance Flags ---
FORCE_CPU = False # Set to True to force CPU usage even if GPU is available (for debugging)

# ===========================================================
# --- Helper Functions ---
# ===========================================================

def clean_extracted_text(text: str | None) -> str:
    """
    Removes null bytes and collapses various whitespace characters into single spaces.

    Args:
        text (str | None): The input text, possibly containing unwanted characters.

    Returns:
        str: The cleaned text.
    """
    if not text:
        return ""
    try:
        # Remove null bytes which can cause issues in various processing steps
        cleaned = text.replace('\x00', '')
        # Replace sequences of whitespace (space, tab, newline, etc.) with a single space
        cleaned = re.sub(r'\s+', ' ', cleaned)
        return cleaned.strip() # Remove leading/trailing whitespace
    except Exception as e:
        print(f"[WARNING] Error during text cleaning: {e}")
        return text # Return original text if cleaning fails


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extracts text from paragraphs and tables in the main body of a DOCX file.
    Includes markers for table content for potential context awareness.

    Args:
        docx_path (str): Path to the DOCX file.

    Returns:
        str: Cleaned, concatenated text content, or empty string on error.
    """
    if not DOCX_AVAILABLE:
        print("[ERROR] Cannot extract DOCX text: python-docx not installed.")
        return ""
    print(f"[INFO] Extracting text from DOCX: {os.path.basename(docx_path)}")
    full_text_parts = []
    try:
        doc = Document(docx_path)
        # Extract from paragraphs
        for para in doc.paragraphs:
            full_text_parts.append(para.text)

        # Extract from tables
        if doc.tables:
            print(f"  [INFO] Extracting text from {len(doc.tables)} table(s) in DOCX.")
            for i, table in enumerate(doc.tables):
                full_text_parts.append(f"\n[Table {i+1} Content Start]") # Marker
                for row_idx, row in enumerate(table.rows):
                    cell_contents = [
                        " ".join([p.text.strip() for p in cell.paragraphs]) # Join paragraphs within cell
                        for cell in row.cells
                    ]
                    # Join cells with a separator, filter out empty cell strings
                    full_text_parts.append(" | ".join(filter(None, cell_contents)))
                full_text_parts.append(f"[Table {i+1} Content End]\n") # Marker

        raw_text = "\n".join(full_text_parts)
        cleaned_text = clean_extracted_text(raw_text)
        print(f"[SUCCESS] DOCX text extracted. Length: {len(cleaned_text)} chars.")
        return cleaned_text
    except PackageNotFoundError:
         print(f"[ERROR] Failed to load DOCX. File might be corrupted or not a valid DOCX: {docx_path}")
         return ""
    except Exception as e:
        print(f"[ERROR] Failed extracting text from DOCX '{os.path.basename(docx_path)}': {e}")
        traceback.print_exc()
        return ""


def extract_images_from_docx(docx_path: str, output_folder: str) -> list[str]:
    """
    Extracts images from DOCX using relationships (preferred) and inline shapes (fallback).
    Saves extracted images as PNG files in the specified output folder.

    Args:
        docx_path (str): Path to the DOCX file.
        output_folder (str): Directory to save extracted images.

    Returns:
        list[str]: A list of file paths to the successfully extracted and saved images.
    """
    if not DOCX_AVAILABLE or not PIL_AVAILABLE:
         print("[WARNING] Cannot extract DOCX images: python-docx or Pillow not available.")
         return []

    print(f"[INFO] Extracting images from DOCX: {os.path.basename(docx_path)}")
    image_paths = []
    processed_rids = set() # Track relationship IDs already processed
    img_counter = 0      # Unique counter for filenames

    # --- Method 1: Relationships (Checks document-part relationships for image types) ---
    # print("  [DEBUG] Attempting DOCX image extraction via relationships...")
    try:
        doc = Document(docx_path)
        doc_part = doc.part
        rels = doc_part.rels

        for rId, rel in rels.items():
            # Check if the relationship target is an image part and not already processed
            if not rel.is_external and hasattr(rel.target_part, 'content_type') \
               and rel.target_part.content_type.startswith('image/') and rId not in processed_rids:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob

                    # Create unique filename and save consistently as PNG
                    safe_partname = f"rel_{rId}"
                    save_ext = 'png'
                    img_filename = f"{TEMP_IMAGE_PREFIX}docx_{img_counter}_{safe_partname}.{save_ext}"
                    img_path = os.path.join(output_folder, img_filename)

                    with Image.open(io.BytesIO(image_bytes)) as img:
                        # Basic conversion for compatibility (e.g., palette to RGBA)
                        if img.mode == 'P' and 'transparency' in img.info: img = img.convert('RGBA')
                        elif img.mode == 'LA': img = img.convert('RGBA')
                        elif img.mode not in ['RGB', 'RGBA']: img = img.convert('RGB') # Convert other modes to RGB
                        img.save(img_path, format=save_ext.upper())
                        image_paths.append(img_path)
                        img_counter += 1
                        processed_rids.add(rId)
                        # print(f"    [DEBUG] Saved image from relationship {rId} to {img_filename}")

                except Exception as img_proc_e:
                    print(f"    [WARNING] Error processing DOCX image relation {rId}: {img_proc_e}")

    except PackageNotFoundError:
         print(f"[ERROR] Failed to load DOCX for image extraction. File might be corrupted.")
         return [] # Cannot proceed
    except Exception as e:
        print(f"[WARNING] Error during DOCX relationship-based image extraction: {e}")
        # Continue to try inline shapes method even if rels had issues

    # print(f"  [DEBUG] Extracted {len(processed_rids)} unique images via relationships.")

    # --- Method 2: Inline Shapes (Checks drawing elements within paragraphs/runs) ---
    # print("  [DEBUG] Attempting DOCX image extraction via inline shapes...")
    inline_img_counter = 0
    try:
        # Ensure doc object exists (might fail above if PackageNotFound)
        if 'doc' not in locals(): doc = Document(docx_path)

        for para in doc.paragraphs:
            for run in para.runs:
                # Use XPath to find inline drawing elements containing an embedded image reference (a:blip)
                if run.element.xpath('.//wp:inline'):
                     for inline in run.element.xpath('.//wp:inline'):
                         blip_fills = inline.xpath('.//a:blip', namespaces=inline.nsmap)
                         if blip_fills:
                              rId = blip_fills[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                              # Process only if this relationship ID hasn't been handled by Method 1
                              if rId and rId not in processed_rids:
                                  try:
                                      image_part = doc.part.related_parts[rId] # Get the image part using the rId
                                      image_bytes = image_part.blob
                                      safe_partname = f"inline_{rId}"
                                      save_ext = 'png'
                                      img_filename = f"{TEMP_IMAGE_PREFIX}docx_{img_counter}_{safe_partname}.{save_ext}"
                                      img_path = os.path.join(output_folder, img_filename)

                                      with Image.open(io.BytesIO(image_bytes)) as img:
                                           # Ensure compatibility before saving
                                           if img.mode == 'P' and 'transparency' in img.info: img = img.convert('RGBA')
                                           elif img.mode == 'LA': img = img.convert('RGBA')
                                           elif img.mode not in ['RGB', 'RGBA']: img = img.convert('RGB')
                                           img.save(img_path, format=save_ext.upper())
                                           image_paths.append(img_path)
                                           img_counter += 1      # Use the main counter
                                           inline_img_counter += 1 # Track inline specifically if needed
                                           processed_rids.add(rId) # Mark as processed
                                           # print(f"    [DEBUG] Saved image from inline shape {rId} to {img_filename}")

                                  except KeyError: continue # rId might not be in related_parts (shouldn't happen often)
                                  except Exception as img_proc_e: print(f"    [WARNING] Error processing inline DOCX image {rId}: {img_proc_e}")
    except Exception as e:
        print(f"[WARNING] Error during DOCX inline image extraction: {e}")

    # print(f"  [DEBUG] Extracted {inline_img_counter} images via inline shapes.")
    total_extracted = len(image_paths)
    if total_extracted > 0:
        print(f"[SUCCESS] Extracted {total_extracted} unique images from DOCX.")
    else:
        print("[INFO] No images found or extracted from DOCX.")
    return image_paths


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extracts text from all pages in a PDF file using PyMuPDF, attempting reading order.

    Args:
        pdf_path (str): Path to the PDF file.

    Returns:
        str: Cleaned, concatenated text content, or empty string on error.
    """
    if not PYMUPDF_AVAILABLE:
        print("[ERROR] Cannot extract PDF text: PyMuPDF not installed.")
        return ""
    print(f"[INFO] Extracting text from PDF: {os.path.basename(pdf_path)}")
    text_parts = []
    try:
        with fitz.open(pdf_path) as doc:
            print(f"  [DEBUG] PDF has {len(doc)} pages.")
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                # Extract text blocks with sorting for better reading order preservation
                blocks = page.get_text("blocks", sort=True)
                # Filter small/insignificant blocks and join text from relevant blocks
                page_text = " ".join([
                    b[4].replace('\n', ' ').strip() # Text content is item 4 in block tuple
                    for b in blocks
                    if abs(b[3] - b[1]) > 5 and abs(b[2] - b[0]) > 5 # Filter blocks smaller than 5x5 pixels
                ])
                text_parts.append(page_text)
            raw_text = "\n".join(text_parts) # Join pages with newline
            cleaned_text = clean_extracted_text(raw_text)
            print(f"[SUCCESS] PDF text extracted. Length: {len(cleaned_text)} chars.")
            return cleaned_text
    except Exception as e:
        print(f"[ERROR] Failed extracting text from PDF '{os.path.basename(pdf_path)}': {e}")
        traceback.print_exc()
        return ""


def extract_images_from_pdf(pdf_path: str, output_folder: str) -> list[str]:
    """
    Extracts images from a PDF file using PyMuPDF, filtering small images.
    Saves extracted images as PNG files in the specified output folder.

    Args:
        pdf_path (str): Path to the PDF file.
        output_folder (str): Directory to save extracted images.

    Returns:
        list[str]: A list of file paths to the successfully extracted and saved images.
    """
    if not PYMUPDF_AVAILABLE or not PIL_AVAILABLE:
         print("[WARNING] Cannot extract PDF images: PyMuPDF or Pillow not available.")
         return []

    print(f"[INFO] Extracting images from PDF: {os.path.basename(pdf_path)}")
    image_paths = []
    img_counter = 0      # Unique counter for filenames
    extracted_xrefs = set() # Track processed image objects by their PDF cross-reference ID

    try:
        doc = fitz.open(pdf_path)
        num_pages = len(doc)
        print(f"  [DEBUG] Processing {num_pages} pages for images...")

        for page_num in range(num_pages):
            try:
                # get_page_images is generally reliable
                img_list = doc.get_page_images(page_num, full=True) # full=True provides more details
                if not img_list: continue # Skip page if no images found

                # print(f"  [DEBUG] Found {len(img_list)} potential image references on PDF page {page_num + 1}")

                for img_index, img_info in enumerate(img_list):
                    xref = img_info[0] # XREF is the unique identifier for the image object within the PDF
                    if xref in extracted_xrefs: continue # Skip if this image object was already processed

                    # Attempt to extract the raw image data
                    base_image = doc.extract_image(xref)
                    if not base_image: # Check if extraction failed (e.g., corrupted image object)
                         # print(f"    [DEBUG] Skipping null image data for xref {xref} on page {page_num+1}")
                         continue

                    # --- Filter based on image dimensions ---
                    img_width = base_image.get("width", 0)
                    img_height = base_image.get("height", 0)
                    if img_width < MIN_IMAGE_DIMENSION_PDF or img_height < MIN_IMAGE_DIMENSION_PDF:
                        # print(f"    [DEBUG] Skipping small image (xref {xref}): {img_width}x{img_height}")
                        continue

                    image_bytes = base_image["image"] # The raw image bytes
                    ext = base_image["ext"]           # File extension detected by PyMuPDF (e.g., 'png', 'jpeg')
                    if not ext: continue              # Skip if format couldn't be determined

                    # --- Save Image Consistently (as PNG using Pillow) ---
                    save_ext = 'png' # Standardize output to PNG for simplicity
                    img_filename = f"{TEMP_IMAGE_PREFIX}pdf_p{page_num+1}_x{xref}_{img_counter}.{save_ext}" # Include page, xref, counter
                    img_path = os.path.join(output_folder, img_filename)

                    try:
                        with Image.open(io.BytesIO(image_bytes)) as img:
                            # Convert modes for better compatibility before saving
                            if img.mode == 'P' and 'transparency' in img.info: img = img.convert('RGBA')
                            elif img.mode == 'LA': img = img.convert('RGBA')
                            elif img.mode not in ['RGB', 'RGBA']: img = img.convert('RGB') # Convert others (like CMYK) to RGB
                            img.save(img_path, format=save_ext.upper())
                            image_paths.append(img_path)
                            extracted_xrefs.add(xref) # Mark this PDF image object as processed
                            img_counter += 1
                            # print(f"    [DEBUG] Saved image (xref {xref}) to {img_filename}")
                    except Exception as img_save_e:
                         print(f"    [WARNING] Error opening/saving extracted PDF image (xref {xref}): {img_save_e}")

            except Exception as page_e:
                 print(f"    [WARNING] Error processing images on PDF page {page_num+1}: {page_e}")
                 # Continue to next page even if one page fails

        doc.close() # Close the PDF document

    except Exception as e:
        print(f"[ERROR] Failed opening or processing PDF for images '{os.path.basename(pdf_path)}': {e}")
        traceback.print_exc()

    if img_counter > 0:
        print(f"[SUCCESS] Extracted {img_counter} unique images (filtered by size) from PDF.")
    else:
        print("[INFO] No images found or extracted from PDF (or they were too small).")
    return image_paths


def check_model_availability(model_name: str) -> bool:
    """
    Checks if the specified Hugging Face model's configuration is accessible.

    Args:
        model_name (str): The name of the model on Hugging Face Hub (e.g., "facebook/bart-large-cnn").

    Returns:
        bool: True if the model config can be downloaded, False otherwise.
    """
    if not TRANSFORMERS_AVAILABLE: return False # Cannot check if library isn't installed
    print(f"[INFO] Checking availability of model configuration: {model_name}")
    try:
        # Attempt to download just the config file as a lightweight check
        hf_hub_download(repo_id=model_name, filename="config.json")
        print(f"[SUCCESS] Model '{model_name}' configuration found on Hub.")
        return True
    except HubNotFoundError:
        print(f"[ERROR] Model '{model_name}' not found on Hugging Face Hub.")
        print("  >>> Suggestion: Verify the SUMMARIZER_MODEL name in config.py is correct.")
        return False
    except requests.exceptions.RequestException as net_err:
        # Catch network errors specifically
        print(f"[ERROR] Network error checking model '{model_name}': {net_err}")
        print("  >>> Suggestion: Check internet connection and firewall settings.")
        return False
    except Exception as check_err:
        # Catch any other errors during the check
        print(f"[ERROR] Unexpected error checking model config for '{model_name}': {check_err}")
        traceback.print_exc()
        return False


def summarize_text(text_to_summarize: str) -> str:
    """
    Summarizes the input text using the configured transformer model via Hugging Face pipeline.

    Handles model loading, tokenization, calculating summary length limits,
    optional manual input truncation, pipeline execution, and error handling.

    Args:
        text_to_summarize (str): The text content extracted from the document.

    Returns:
        str: The generated summary text, or a user-friendly error message string if summarization failed.
    """
    if not TRANSFORMERS_AVAILABLE:
        return "[ERROR: Summarization libraries (transformers/torch) not installed]"
    if not text_to_summarize or not text_to_summarize.strip():
        print("[WARNING] No text provided or extracted for summarization.")
        return "No content available to summarize."

    # Optional: Check model availability early - uncomment if desired
    # if not check_model_availability(SUMMARIZER_MODEL):
    #     return f"Error: Summarization model '{SUMMARIZER_MODEL}' is not accessible."

    # --- Initialize Pipeline ---
    print(f"\n[INFO] Initializing summarization pipeline with model: {SUMMARIZER_MODEL}")
    summarizer = None
    try:
        # Determine device: Use GPU if available and not forced to CPU
        pipeline_device = 0 if torch.cuda.is_available() and not FORCE_CPU else -1
        device_name = "GPU:0" if pipeline_device == 0 else "CPU"
        if FORCE_CPU and torch.cuda.is_available(): print("  [DEBUG] FORCE_CPU flag is True. Using CPU despite GPU availability.")

        start_time = time.time()
        summarizer = transformers.pipeline(
            "summarization",
            model=SUMMARIZER_MODEL,
            device=pipeline_device,
            framework="pt" # Explicitly use PyTorch framework
            # You might add revision='main' if needed for specific model versions
        )
        load_time = time.time() - start_time
        print(f"[SUCCESS] Summarizer pipeline loaded in {load_time:.2f}s. Using device: {device_name}")
    except ImportError as ie:
         # Catch errors if a dependency is missing despite transformers being found
         print(f"[FATAL ERROR] Missing dependency for transformers pipeline: {ie}")
         print("  >>> Suggestion: Ensure 'torch' (or 'tensorflow') and potentially 'accelerate' are installed correctly.")
         traceback.print_exc()
         return "Error: Missing dependency for summarization pipeline."
    except Exception as e:
        print(f"[FATAL ERROR] Could not initialize pipeline for model '{SUMMARIZER_MODEL}'. Error: {e}")
        if 'init_empty_weights' in str(e) or 'accelerate' in str(e):
            print("  >>> Suggestion: This often indicates a library version mismatch or missing/incompatible 'accelerate'.")
            print("  >>> Try: pip install accelerate --upgrade")
            print("  >>> Also try: pip install transformers torch --upgrade")
        traceback.print_exc()
        return "Error: Summarization pipeline initialization failed (check logs)."

    # --- Tokenize Input and Calculate Lengths ---
    processed_text = text_to_summarize # Text to actually feed into the pipeline
    num_tokens = 0
    model_max_len = 512 # Default fallback, updated by tokenizer
    min_summary_len = SUMMARY_ABS_MIN_LENGTH # Initialize with defaults
    max_summary_len = SUMMARY_ABS_MAX_LENGTH # Initialize with defaults

    try:
        print("[INFO] Tokenizing input text for length calculation...")
        tokenizer = summarizer.tokenizer # Get tokenizer from the loaded pipeline
        # Encode without truncation first to get the true input token count
        tokens = tokenizer(text_to_summarize, return_tensors="pt", truncation=False)
        num_tokens = tokens.input_ids.shape[1]
        # Get the model's maximum input length from the tokenizer config
        model_max_len = tokenizer.model_max_length
        print(f"  [INFO] Original input text length: {num_tokens} tokens.")
        print(f"  [INFO] Model max sequence length (from tokenizer): {model_max_len}")

        # --- Manual Input Truncation (Optional - based on config flag) ---
        if MANUAL_TRUNCATE_INPUT and num_tokens > MANUAL_TRUNCATE_TOKEN_LIMIT:
            print(f"[INFO] Input ({num_tokens} tokens) exceeds manual limit ({MANUAL_TRUNCATE_TOKEN_LIMIT}). Truncating input...")
            # Truncate the token IDs tensor
            # [0] accesses the first (and only) batch item, :limit slices tokens
            truncated_ids = tokens.input_ids[0, :MANUAL_TRUNCATE_TOKEN_LIMIT].unsqueeze(0) # Add batch dim back
            # Decode the truncated IDs back into text - use skip_special_tokens
            processed_text = tokenizer.decode(truncated_ids[0], skip_special_tokens=True, clean_up_tokenization_spaces=True)
            original_length = num_tokens
            # Recalculate token count *after* decoding, as it might slightly change
            num_tokens = tokenizer(processed_text, return_tensors="pt", truncation=False).input_ids.shape[1]
            print(f"  [INFO] Input manually truncated to {num_tokens} tokens (from {original_length}).")
            # print(f"  [DEBUG] Manually truncated text sample: {processed_text[:500]}...") # For debugging
        elif num_tokens > model_max_len:
             # Log a warning if the input still exceeds the model's max length,
             # indicating the pipeline's internal truncation will be used.
             print(f"  [WARNING] Input ({num_tokens} tokens) exceeds model max length ({model_max_len}). Pipeline's internal truncation will be used.")
        # --- End Manual Truncation ---

        # --- Calculate Summary Length Limits ---
        # Base calculation on the (potentially truncated) input token count
        max_summary_len = min(SUMMARY_ABS_MAX_LENGTH, int(num_tokens * SUMMARY_MAX_LENGTH_FACTOR))
        min_summary_len = max(SUMMARY_ABS_MIN_LENGTH, int(num_tokens * SUMMARY_MIN_LENGTH_FACTOR))
        # Ensure min < max and apply absolute bounds and reasonable gap
        max_summary_len = max(max_summary_len, min_summary_len + 10) # Max should be at least min+10
        min_summary_len = min(min_summary_len, max(SUMMARY_ABS_MIN_LENGTH, max_summary_len - 10)) # Min shouldn't exceed max-10
        min_summary_len = max(min_summary_len, 10) # Absolute floor for min length
        # Optional: Cap max_summary_len at the model's max input length if desired
        # max_summary_len = min(max_summary_len, model_max_len)
        print(f"  [INFO] Target summary length (tokens): min={min_summary_len}, max={max_summary_len}")

    except Exception as tok_e:
        print(f"[WARNING] Tokenizer error occurred ({tok_e}). Could not determine precise token count or model max length.")
        traceback.print_exc()
        # Fallback to using absolute limits only if tokenization fails
        min_summary_len = SUMMARY_ABS_MIN_LENGTH
        max_summary_len = SUMMARY_ABS_MAX_LENGTH
        print(f"  [WARNING] Using absolute summary length limits due to tokenizer error: min={min_summary_len}, max={max_summary_len}")
        # processed_text remains the original cleaned text in this fallback


    # --- Execute Summarization Pipeline ---
    try:
        print("[INFO] Starting summarization process...")
        start_time = time.time()
        # Feed the 'processed_text' (which might be manually truncated) to the pipeline
        # The pipeline's truncation=True handles cases where input *still* exceeds model_max_len
        summary_result = summarizer(
            processed_text,
            max_length=max_summary_len,
            min_length=min_summary_len,
            no_repeat_ngram_size=3, # Helps reduce repetitive phrases like "is is is"
            # early_stopping=True,  # Can stop generation early if end condition met
            do_sample=False,        # Use deterministic greedy search (usually better for summaries)
            truncation=True         # Crucial: Allows pipeline to handle inputs longer than model max length internally
        )
        summarize_time = time.time() - start_time

        # --- Validate and Extract Summary Text ---
        if not summary_result or not isinstance(summary_result, list) or not summary_result[0].get('summary_text'):
            print(f"[ERROR] Summarization pipeline returned an unexpected result format: {summary_result}")
            return "Error: Summarization failed to produce expected output format."

        summary_text = summary_result[0]['summary_text']
        summary_char_len = len(summary_text)
        # Optional: Calculate summary token length if needed for logging
        # summary_token_len = tokenizer(summary_text, return_tensors="pt").input_ids.shape[1]
        print(f"[SUCCESS] Summarization complete in {summarize_time:.2f}s.")
        print(f"  [INFO] Generated summary length: {summary_char_len} characters.") # More intuitive than tokens for user
        return summary_text

    except IndexError:
        # Should be caught by the validation above, but included as a failsafe.
        print(f"[ERROR] Summarization pipeline returned an empty list or unexpected structure.")
        return "Error: Summarization failed to produce output (IndexError)."
    except Exception as e:
        print(f"[ERROR] An error occurred during summarization execution: {e}")
        # Provide specific suggestions for common errors
        err_str = str(e).upper()
        if "CUDA" in err_str or "OUT OF MEMORY" in err_str:
            print("  >>> CUDA Error encountered. Possible causes:")
            print("      - Input text too long/complex for GPU memory, even with truncation.")
            print("      - CUDA driver / PyTorch / Transformers / Accelerate version incompatibility.")
            print("      - Insufficient GPU memory or other processes using GPU.")
            print("  >>> Suggestions:")
            print("      - Reduce MANUAL_TRUNCATE_TOKEN_LIMIT further if enabled.")
            print("      - Set FORCE_CPU=True in report_summarizer.py to test on CPU (slower).")
            print("      - Try a smaller summarization model (e.g., distilbart).")
            print("      - Check GPU utilization and memory (e.g., using `nvidia-smi`).")
            # Suggest setting env var for detailed CUDA errors *before* running script:
            # export CUDA_LAUNCH_BLOCKING=1 (Linux/macOS) or set CUDA_LAUNCH_BLOCKING=1 (Windows)
        elif "maximum sequence length" in str(e).lower():
             print("  >>> Error likely due to sequence length exceeding model limits.")
             print("     (This shouldn't happen often with truncation=True in pipeline).")
             print("  >>> Suggestion: Ensure manual truncation limit is well below model max length.")

        traceback.print_exc() # Print full traceback for debugging
        return f"Error during summarization execution (check logs)."


def create_summary_document(summary_text: str, image_paths: list[str]) -> Document | None:
    """
    Creates a new DOCX document containing the summary text and extracted images.

    Args:
        summary_text (str): The generated summary (or an error message string).
        image_paths (list[str]): List of local file paths to the extracted images.

    Returns:
        docx.Document | None: A new python-docx Document object, or None if DOCX library is unavailable.
    """
    if not DOCX_AVAILABLE:
        print("[ERROR] Cannot create summary document: python-docx not available.")
        return None

    print("[INFO] Creating summary DOCX document structure...")
    doc = Document()
    added_image_count = 0
    try:
        # --- Add Summary Section ---
        doc.add_heading('Generated Summary', level=1)
        # Add the summary text (could be actual summary or an error message from summarize_text)
        para = doc.add_paragraph(summary_text if summary_text else "[No summary generated or an error occurred]")
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # Justify text alignment
        doc.add_paragraph() # Add some vertical space after the summary

        # --- Add Images Section ---
        if image_paths:
            # Filter list to include only paths that currently exist
            valid_image_paths = [p for p in image_paths if p and os.path.exists(p)]
            if valid_image_paths:
                print(f"[INFO] Adding {len(valid_image_paths)} extracted images to the document...")
                # Add page break and heading only if there are valid images
                try: doc.add_page_break() # Add break before images section
                except Exception as pb_e: print(f"[WARNING] Could not add page break: {pb_e}") # Non-critical error
                doc.add_heading('Extracted Images', level=1)

                # Add each valid image, centered
                for idx, img_path in enumerate(valid_image_paths):
                    # print(f"  [DEBUG] Adding image {idx + 1}/{len(valid_image_paths)}: {os.path.basename(img_path)}")
                    try:
                        # Create a new paragraph for each image for better spacing control
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # Center the image
                        run = img_para.add_run()
                        # Add the picture with specified width
                        run.add_picture(img_path, width=Inches(IMAGE_WIDTH_INCHES_SUMMARY))
                        added_image_count += 1
                    except FileNotFoundError:
                         print(f"    [ERROR] Image file not found during add_picture: {img_path}")
                    except Exception as e:
                         print(f"    [ERROR] Failed to add image '{os.path.basename(img_path)}' to document: {e}")
                         if "unsupported image format" in str(e).lower():
                             print("      >>> Suggestion: Check if Pillow conversion to PNG worked correctly.")
                print(f"[INFO] Finished adding images. Successfully added {added_image_count}/{len(valid_image_paths)}.")
            else:
                print("[INFO] No valid extracted image files found to add after checking paths.")
        else:
            print("[INFO] No image paths were provided for inclusion in the summary document.")

        print("[SUCCESS] Summary document structure created.")
        return doc
    except Exception as e:
         print(f"[ERROR] Failed to create the overall summary document structure: {e}")
         traceback.print_exc()
         # Return None if document creation fails critically
         return None

# ===========================================================
# --- Main Report Generation Function (Entry Point) ---
# ===========================================================
def generate_report(data: dict, assets: dict) -> tuple[Document | None, list[str]]:
    """
    Main function for the Summarizer report type.

    Orchestrates the extraction of text/images from the uploaded file,
    calls the summarization pipeline, and creates the final DOCX output document.

    Args:
        data (dict): Dictionary of sheet data for the current row (may be used for context).
        assets (dict): Dictionary containing the path to the uploaded report file.
                       Expected key: 'report_to_summarize'.

    Returns:
        tuple[Document | None, list[str]]:
            - The generated python-docx Document object (or None on critical failure).
            - A list of paths to temporary extracted image files created during the process
              (used for cleanup by the main script).
    """
    start_time_summary = time.time()
    print("\n=== Starting Summary Report Generation ===")
    uploaded_report_path = assets.get('report_to_summarize')
    generated_doc = None           # Initialize doc object
    extracted_image_paths = []     # Initialize list to store paths of created temp images

    # --- Input Validation ---
    if not uploaded_report_path or not os.path.exists(uploaded_report_path):
        print(f"[ERROR] Input report file not found at expected path: {uploaded_report_path}")
        return None, [] # Return None doc, empty image list

    print(f"[INFO] Input file for summarization: {os.path.basename(uploaded_report_path)}")
    _, file_extension = os.path.splitext(uploaded_report_path)
    file_extension = file_extension.lower()

    # --- Content Extraction ---
    extracted_text = ""
    temp_folder_for_images = config.TEMP_FOLDER # Use globally configured temp folder

    if file_extension == '.pdf':
        print("[INFO] Detected PDF file type. Starting extraction...")
        extracted_text = extract_text_from_pdf(uploaded_report_path)
        extracted_image_paths = extract_images_from_pdf(uploaded_report_path, temp_folder_for_images)
    elif file_extension == '.docx':
        print("[INFO] Detected DOCX file type. Starting extraction...")
        extracted_text = extract_text_from_docx(uploaded_report_path)
        extracted_image_paths = extract_images_from_docx(uploaded_report_path, temp_folder_for_images)
    else:
        print(f"[ERROR] Unsupported file type '{file_extension}'. Only .pdf and .docx are supported for summarization.")
        return None, [] # Cannot process, return None doc, empty image list

    # --- Summarization ---
    if not extracted_text.strip():
         print("[WARNING] No text could be extracted from the document. Summary will reflect this.")
         # Set summary text to indicate no content, rather than calling summarize_text
         summary_text = "No text content could be extracted from the uploaded document."
    else:
         print("[INFO] Text extracted successfully. Proceeding to summarization...")
         # Call the summarization function (which handles internal errors and returns message on failure)
         summary_text = summarize_text(extracted_text)

    # --- Document Creation ---
    # Always attempt to create the document, even if summary_text contains an error message.
    # This provides the user with feedback and any extracted images.
    generated_doc = create_summary_document(summary_text, extracted_image_paths)
    if generated_doc is None:
         print("[ERROR] Final summary document object creation failed.")
         # Return None, but keep image paths if they were extracted, for potential cleanup
         return None, extracted_image_paths

    end_time_summary = time.time()
    print(f"=== Summary Report Generation Finished ({end_time_summary - start_time_summary:.2f}s) ===")

    # Return the generated document object AND the list of temporary image paths created
    return generated_doc, extracted_image_paths